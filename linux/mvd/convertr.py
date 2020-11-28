import os
import re
import docx
import pymorphy2
from docx import Document
from odf import text, teletype
from odf.opendocument import load, OpenDocumentText
from tkinter import *
from tkinter.filedialog import askopenfilename, askdirectory
from tkinter import messagebox as mb

morph = pymorphy2.MorphAnalyzer()
Dictionary_After1_Error = {"Ко нему": "К нему", "ко нему": "к нему", "Ко ней": "К ней", "ко ней": "к ней",
                           "Ко ним": "К ним", "ко ним": "к ним",
                           "Ко ему": "К нему", "ко ему": "к нему", "Ко ей": "К ней", "ко ей": "к ней",
                           "Ко им": "К ним", "ко им": "к ним", "К ей": "К ней", "к ей": "к ней",
                           "У их": "У них", "у их": "у них",
                           }

def conv():
    global filepath
    global ch_box
    filename = os.path.split(filepath)
    if filename[1].endswith(".docx") is True or filename[1].endswith(".odt") is True:
        filepath2 = filename[0] + "/" + "Измененный_" + filename[1]

        if filename[1].endswith(".docx"):
            file = Document(filepath)
            if os.path.isfile(filepath2) is False:
                document = docx.Document()
                document.save("Измененный_" + filename[1])
            file2 = Document(filepath2)

            gender = file.paragraphs[13].text.split()
            if gender[4].endswith(('а', 'я')):
                gender = 'femn'
            else:
                gender = 'masc'
            Str = 52
            Flag = '.docx'
        else:
            file = load(filepath)
            if os.path.isfile(filepath2) is False:
                document = OpenDocumentText()
                document.save("Измененный_" + filename[1])
            file2 = load(filepath2)
            allparas = file.getElementsByType(text.P)

            gender = teletype.extractText(allparas[14]).split()
            if gender[4].endswith(('а', 'я')):
                gender = 'femn'
            else:
                gender = 'masc'
            Str = 55
            Flag = '.odt'
    else:
        # try:
        filedirectory = askdirectory()
        # except :
        #     return
        # else:
        if filedirectory == '':
            return
        else:
            filename = "Конвертированный файл.txt"
            file = open(filedirectory + "/" + filename, 'w')

            if sex.get() == 0:
                gender = 'masc'
            else:
                gender = 'femn'

            Flag = 'textbox'
    if Flag != 'textbox':
        Count = 0
    else:
        Count = 1

    while Count != 2:
        if Flag == '.docx':
            if file.paragraphs[Str].text == '':
                Count += 1
                Str += 1
                continue
            else:
                offers = re.split(r'\.', file.paragraphs[Str].text)
        elif Flag == '.odt':
            if teletype.extractText(allparas[Str]) == '':
                Count += 1
                Str += 1
                continue
            else:
                offers = re.split(r'\.', teletype.extractText(allparas[Str]))
        else:
            offers = re.split(r'\.', my_text.get(1.0, END))

        textline = ''
        for k in range(0, len(offers)):
            string = re.findall(r'\w+', offers[k])
            for i in string:
                wd = ''
                word = morph.parse(i)[0]
                if i in ["зашли", "стоит"]:
                    word = morph.parse(i)[1]
                if word.tag.POS == 'NPRO' and (word.tag.person == '1per' or word.tag.person == '2per'):
                    if i in ["Я", "я", "Меня", "меня", "Мне", "мне", "Мной", "мной", "Ты", "ты", "Тебя", "тебя",
                             "Тебе", "тебе", "Тобой", "тобой"]:
                        if gender == 'masc':
                            wd = morph.parse("Он")[0]
                        else:
                            wd = morph.parse("Она")[0]
                    else:
                        if i in ["Мы", "мы", "Нас", "нас", "Нам", "нам", "Нами", "нами", "Вы", "вы", "Вас", "вас",
                                 "Вам", "вам", "Вами", "вами"]:
                            wd = morph.parse("Они")[0]
                    wd = wd.inflect({word.tag.case})[0]
                    if i.istitle() is True:
                        wd = wd.capitalize()
                    offers[k] = re.sub(r'\b' + i + r'\b', wd, offers[k], count=1)
                if word.tag.POS == 'ADJF' and 'Apro' in word.tag:
                    list1 = ["мой", "моего", "моему", "моим", "моём", "моем", "твой", "твоего", "твоему", "твоим",
                             "твоём", "твоем", "моя", "моей", "мою", "твоя", "твоей", "твою", "моё", "мое",
                             "мои", "моих", "моим", "моими", "твои", "твоих", "твоим", "твоими"]

                    list2 = ["наш", "нашего", "нашему", "нашим", "нашем", "ваш", "вашего", "вашему", "вашим",
                             "вашем", "наше", "ваше", "наша", "нашей", "нашу", "ваша", "вашей", "вашу",
                             "наши", "наших", "нашим", "нашими", "ваши", "ваших", "вашим", "вашими"]
                    for j in list1:
                        if i == j or i == j.capitalize():
                            if gender == 'masc':
                                wd = 'его'
                            else:
                                wd = 'её'
                            if i.istitle() is True:
                                wd = wd.capitalize()
                            offers[k] = re.sub(r'\b' + i + r'\b', wd, offers[k], count=1)
                            break
                    for j in list2:
                        if i == j or i == j.capitalize():
                            wd = 'их'
                            if i.istitle() is True:
                                wd = wd.capitalize()
                            offers[k] = re.sub(r'\b' + i + r'\b', wd, offers[k], count=1)
                            break
                if {'VERB', '1per', 'sing', 'pres'} in word.tag:
                    wd = word.inflect({'VERB', 'sing', 'pres', '3per'})[0]
                    if i.istitle() is True:
                        wd = wd.capitalize()
                    offers[k] = re.sub(r'\b' + i + r'\b', wd, offers[k], count=1)
            for i in Dictionary_After1_Error:
                offers[k] = re.sub(str(i), Dictionary_After1_Error[i], offers[k])
        for i in offers:
            if i != '':
                if i.endswith('!') or i.endswith('?'):
                    textline += i
                else:
                    textline += i + '.'
            else:
                continue
        if Flag == ".docx":
            file2.add_paragraph(textline, None)
        elif Flag == ".odt":
            file2.text.addElement(text.P(text=textline))
        else:
            file.write(textline)
            Count += 1
        if Flag != "textbox":
            Str += 1
    if Flag != "textbox":
        file2.save(filepath2)
    else:
        file.close()
    #if ch_box == 1:
    mb.showinfo(title="ВНИМАНИЕ", message="Преобразование окончено!!!" + "\r\n"
                                          + "Преобразованный файл лежит рядом с исходным")
    #else:
        #mb.showinfo(title="ВНИМАНИЕ", message="Преобразование окончено!!!")
    # сюда бы прикрутить открывание папки, но tkinter умеет открывать только для загрузки
    # filepath = askopenfilename(initialdir=".")

#загрузка файла
def add_file():
    global filepath
    try:
        filepath = askopenfilename(initialdir=".", filetypes=(("Docx File", "*.docx"), ("Open Office File", "*.odt"), ("All Files", "*.*")), title="Выберите файл")
    except FileNotFoundError:
        return
    else:
        my_text.delete(1.0, END)
        if filepath.endswith(".docx"):
            file = Document(filepath)
            Str = 52

            Count = 0
            while Count != 2:
                if file.paragraphs[Str].text == '':
                    Count += 1
                    Str += 1
                    continue
                else:
                    my_text.insert(END, file.paragraphs[Str].text + '\n')
                    Str += 1
        else:
            file = load(filepath)
            Str = 55
            allparas = file.getElementsByType(text.P)

            Count = 0
            while Count != 2:
                if teletype.extractText(allparas[Str]) == '':
                    Count += 1
                    Str += 1
                    continue
                else:
                    my_text.insert(END, teletype.extractText(allparas[Str]))
                    Str += 1

def copy_text(event):
    global selected
    if event:
       selected = window.clipboard_get()
    if my_text.selection_get():
        selected = my_text.selection_get()
        window.clipboard_clear()
        window.clipboard_append(selected)

def paste_text(event):
    global selected
    if event:
        selected = window.clipboard_get()
    else:
        if selected:
            position = my_text.index(INSERT)
            my_text.insert(position, selected)

def select_all_text(event):
    my_text.tag_add('sel', '1.0', 'end')

def cut_text(event):
    global selected
    if event:
        selected = window.clipboard_get()
    else:
        if my_text.selection_get():
            selected = my_text.selection_get()
            my_text.delete("sel.first", "sel.last")
            window.clipboard_clear()
            window.clipboard_append(selected)

def clear_text():
    global filepath
    my_text.delete(1.0, END)
    filepath = ""
    sex.set(0)

def clear_text2():
    global filepath
    my_text2.delete(1.0, END)
    filepath = ""
    sex.set(0)

def exit_conv(event):
    #тут надо чекбокс
    answer = mb.askyesno(title="Выход", message="Вы действительно хотите выйти из программы??")
    if answer:
        window.destroy()
    #sys.exit() не работает но пусть пока будет

def out_file():
    pass

#Полезные константы
filepath = ""
selected = ""

# основные настройки
window = Tk()
window.title("Конвертер")
window.minsize(1200, 700)
#принудительно запретить изменение размера окна
#window.resizable(width=False, height=False)
# window.iconphoto(False, PhotoImage(file="F:\Jija Lichnoe\Python\Проекты\Разные штуки\CyberBreak2020_Final_MVD\icon.png"))
# радиобокс для выбора гендера
sex = IntVar()  # храним как int, можно как bool, тогда меняем на BooleanVar()
sex.set(0)  # по умолчанию будет мужской ибо value = 0 у мужского рода и в sex.set в скобках 0
male = Radiobutton(text="Мужской род", variable=sex, value=0)

female = Radiobutton(text="Женский род", variable=sex, value=1)

print_file = BooleanVar()
print_file.set(0)
ch_box = Checkbutton(window, text="Выгружать результат в файл",
                 variable=print_file,
                 onvalue=1, offvalue=0,
                 command=out_file)

# меню для файла
menu = Menu(window)
new_item = Menu(menu)
new_item.add_command(label="Загрузить файл", command=add_file)
menu.add_cascade(label='Файл', menu=new_item)
# меню для редактирования строк
new_item = Menu(menu)
new_item.add_command(label="Копировать", command=lambda: copy_text(0), accelerator="(Ctrl+C)")
new_item.add_command(label="Вставить", command=lambda: paste_text(0), accelerator="(Ctrl+V)")
new_item.add_command(label="Вырезать", command=lambda: cut_text(0), accelerator="(Ctrl+X)")
new_item.add_command(label="Выделить всё", command=lambda: select_all_text(0), accelerator="(Ctrl+A)")
menu.add_cascade(label='Редактировать', menu=new_item)
window.config(menu=menu)
# поясняшка
lbl = Label(window, text="Вставьте в поле ввода необходимый фрагмент текста или загрузите файл и нажмите ПРЕОБРАЗОВАТЬ", font="16")
lbl.place(relx=0.12, rely = 0.09)

lbl1 = Label(window, text = "Поле ввода")
lbl2 = Label(window, text = "Поле вывода")
lbl1.place(relx=0.05, rely = 0.15)
lbl2.place(relx=0.5, rely = 0.15)
# текстовое поле
my_text = Text(window, width=65, height=30)
my_text.pack(pady=5, padx=5)
my_text.place(relx = 0.05, rely =0.2)

my_text2 = Text(window, width=65, height=30)
my_text2.pack(pady=5, padx=5)
my_text2.place(relx=0.50, rely =0.2)

# кнопушки
btn_clear = Button(window, text="Очистить поле ввода", command=clear_text)
btn_conv = Button(window, text="ПРЕОБРАЗОВАТЬ", command=conv)
btn_upload = Button(window, text="Загрузить файл", command=add_file)
btn_clear2 = Button(window, text="Очистить поле вывода", command=clear_text2)

#выравнивание кнопушек
btn_clear.grid(row = 0, column = 0)
btn_clear2.grid(row = 0, column = 1)
btn_upload.grid(row = 0, column = 2)
btn_conv.grid(row = 0, column = 3)
male.grid(row =0, column = 4)
female.grid(row =0, column = 5)
ch_box.grid(row =0, column = 6)


#горячие клавиши
window.bind("<Control-Key-c>", copy_text)
window.bind("<Control-Key-v>", paste_text)
window.bind("<Control-Key-x>", cut_text)
window.bind("<Control-Key-a>", select_all_text)
window.bind("<Escape>", exit_conv)

window.mainloop()
