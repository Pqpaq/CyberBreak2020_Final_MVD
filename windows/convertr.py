import os
import re
import docx
import pymorphy2
from docx import Document
from odf import text, teletype
from odf.opendocument import load, OpenDocumentText
from tkinter import *
from tkinter.filedialog import askopenfilename, askdirectory, asksaveasfilename
from tkinter import messagebox as mb

morph = pymorphy2.MorphAnalyzer()
Dictionary_After1_Error = {"Ко нему": "К нему", "ко нему": "к нему", "Ко ней": "К ней", "ко ней": "к ней",
                           "Ко ним": "К ним", "ко ним": "к ним",
                           "Ко ему": "К нему", "ко ему": "к нему", "Ко ей": "К ней", "ко ей": "к ней",
                           "Ко им": "К ним", "ко им": "к ним", "К ей": "К ней", "к ей": "к ней",
                           "У их": "У них", "у их": "у них", "При ей": "При ней", "при ей": "при ней",
                           "При ему": "При нем", "при ему": "при нем",
                           }

def conv():
    global filepath
    filename = os.path.split(filepath)
    if print_file.get() == 1:
        if filename[1].endswith(".docx") is True or filename[1].endswith(".odt") is True:
            filepath2 = filename[0] + "/" + "Измененный_" + filename[1]

            if filename[1].endswith(".docx"):
                file = Document(filepath)
                if os.path.isfile(filepath2) is False:
                    document = docx.Document()
                    document.save("Измененный_" + filename[1])
                file2 = Document(filepath2)

                gender = file.paragraphs[13].text.split()
                if gender[4].endswith(('а', 'я')):
                    gender = 'femn'
                else:
                    gender = 'masc'
                Str = 52
                Flag = '.docx'
            else:
                file = load(filepath)
                if os.path.isfile(filepath2) is False:
                    document = OpenDocumentText()
                    document.save("Измененный_" + filename[1])
                file2 = load(filepath2)
                allparas = file.getElementsByType(text.P)

                gender = teletype.extractText(allparas[14]).split()
                if gender[4].endswith(('а', 'я')):
                    gender = 'femn'
                else:
                    gender = 'masc'
                Str = 55
                Flag = '.odt'
        else:
            filedirectory = askdirectory()
            if filedirectory == '':
                return
            else:
                filename = "Конвертированный файл.txt"
                file = open(filedirectory + "/" + filename, 'w')

                if sex.get() == 0:
                    gender = 'masc'
                else:
                    gender = 'femn'

                Flag = 'textbox'
        if Flag != 'textbox':
            Count = 0
        else:
            Count = 1

        while Count != 2:
            if Flag == '.docx':
                if file.paragraphs[Str].text == '':
                    Count += 1
                    Str += 1
                    continue
                else:
                    offers = re.split(r'\.', file.paragraphs[Str].text)
            elif Flag == '.odt':
                if teletype.extractText(allparas[Str]) == '':
                    Count += 1
                    Str += 1
                    continue
                else:
                    offers = re.split(r'\.', teletype.extractText(allparas[Str]))
            else:
                offers = re.split(r'\.', my_text.get(1.0, END))

            textline = ''
            for k in range(0, len(offers)):
                stroka = ""
                if offers[k].startswith("-"):
                    continue
                if re.findall(r'(«.{1,}»)', offers[k]):
                    kav = re.findall(r'(«[^«»]{1,}»)', offers[k])
                    for i in kav:
                        stroka += i
                    stroka = stroka.split()
                elif re.findall(r'(".{1,}")', offers[k]):
                    kav = re.findall(r'("[^"]{1,}")', offers[k])
                    for i in kav:
                        stroka += i
                    stroka = stroka.split()
                string = re.findall(r'\w+', offers[k])
                for i in string:
                    if i.startswith('\"') or i.startswith('«') or i in stroka:
                        continue
                    wd = ''
                    word = morph.parse(i)[0]
                    if i in ["зашли", "стоит"]:
                        word = morph.parse(i)[1]
                    if word.tag.POS == 'NPRO' and (word.tag.person == '1per' or word.tag.person == '2per'):
                        if i in ["Я", "я", "Меня", "меня", "Мне", "мне", "Мной", "мной", "Ты", "ты", "Тебя", "тебя",
                                 "Тебе", "тебе", "Тобой", "тобой"]:
                            if gender == 'masc':
                                wd = morph.parse("Он")[0]
                            else:
                                wd = morph.parse("Она")[0]
                        else:
                            if i in ["Мы", "мы", "Нас", "нас", "Нам", "нам", "Нами", "нами", "Вы", "вы", "Вас", "вас",
                                     "Вам", "вам", "Вами", "вами"]:
                                wd = morph.parse("Они")[0]
                        wd = wd.inflect({word.tag.case})[0]
                        if i.istitle() is True:
                            wd = wd.capitalize()
                        offers[k] = re.sub(r'\b' + i + r'\b', wd, offers[k], count=1)
                    if word.tag.POS == 'ADJF' and 'Apro' in word.tag:
                        list1 = ["мой", "моего", "моему", "моим", "моём", "моем", "твой", "твоего", "твоему", "твоим",
                                 "твоём", "твоем", "моя", "моей", "мою", "твоя", "твоей", "твою", "моё", "мое",
                                 "мои", "моих", "моим", "моими", "твои", "твоих", "твоим", "твоими"]

                        list2 = ["наш", "нашего", "нашему", "нашим", "нашем", "ваш", "вашего", "вашему", "вашим",
                                 "вашем", "наше", "ваше", "наша", "нашей", "нашу", "ваша", "вашей", "вашу",
                                 "наши", "наших", "нашим", "нашими", "ваши", "ваших", "вашим", "вашими"]
                        for j in list1:
                            if i == j or i == j.capitalize():
                                if gender == 'masc':
                                    wd = 'его'
                                else:
                                    wd = 'её'
                                if i.istitle() is True:
                                    wd = wd.capitalize()
                                offers[k] = re.sub(r'\b' + i + r'\b', wd, offers[k], count=1)
                                break
                        for j in list2:
                            if i == j or i == j.capitalize():
                                wd = 'их'
                                if i.istitle() is True:
                                    wd = wd.capitalize()
                                offers[k] = re.sub(r'\b' + i + r'\b', wd, offers[k], count=1)
                                break
                    if {'VERB', '1per', 'sing', 'pres'} in word.tag:
                        wd = word.inflect({'VERB', 'sing', 'pres', '3per'})[0]
                        if i.istitle() is True:
                            wd = wd.capitalize()
                        offers[k] = re.sub(r'\b' + i + r'\b', wd, offers[k], count=1)
                for i in Dictionary_After1_Error:
                    offers[k] = re.sub(str(i), Dictionary_After1_Error[i], offers[k])
            for i in offers:
                if i != '':
                    if i.endswith('!') or i.endswith('?'):
                        textline += i
                    else:
                        textline += i + '.'
                else:
                    continue
            if Flag == ".docx":
                file2.add_paragraph(textline, None)
            elif Flag == ".odt":
                file2.text.addElement(text.P(text=textline))
            else:
                file.write(textline + '\n')
                Count += 1
            if Flag != "textbox":
                Str += 1
        if Flag != "textbox":
            file2.save(filepath2)
        else:
            file.close()
        mb.showinfo(title="ВНИМАНИЕ", message="Преобразование окончено!!!" + "\r\n"
                                                + "Преобразованный файл лежит рядом с исходным")
    else:
        if sex.get() == 0:
            gender = 'masc'
        else:
            gender = 'femn'

        Count = 1
        while Count != 2:
            offers = re.split(r'\.', my_text.get(1.0, END))
            textline = ''
            for k in range(0, len(offers)):
                stroka = ""
                if offers[k].startswith("-") or offers[k].startswith("\n-"):
                    continue
                if re.findall(r'(«.{1,}»)', offers[k]):
                    kav = re.findall(r'(«[^«»]{1,}»)', offers[k])
                    for i in kav:
                        stroka += i
                    stroka = stroka.split()
                elif re.findall(r'(".{1,}")', offers[k]):
                    kav = re.findall(r'("[^"]{1,}")', offers[k])
                    for i in kav:
                        stroka += i
                    stroka = stroka.split()
                string = re.findall(r'\w+', offers[k])
                for i in string:
                    if i.startswith('\"') or i.startswith('«') or i.startswith('\n') or i in stroka:
                        continue
                    wd = ''
                    word = morph.parse(i)[0]
                    if i in ["зашли", "стоит"]:
                        word = morph.parse(i)[1]
                    if word.tag.POS == 'NPRO' and (word.tag.person == '1per' or word.tag.person == '2per'):
                        if i in ["Я", "я", "Меня", "меня", "Мне", "мне", "Мной", "мной", "Ты", "ты", "Тебя", "тебя",
                                 "Тебе", "тебе", "Тобой", "тобой"]:
                            if gender == 'masc':
                                wd = morph.parse("Он")[0]
                            else:
                                wd = morph.parse("Она")[0]
                        else:
                            if i in ["Мы", "мы", "Нас", "нас", "Нам", "нам", "Нами", "нами", "Вы", "вы", "Вас", "вас",
                                     "Вам", "вам", "Вами", "вами"]:
                                wd = morph.parse("Они")[0]
                        wd = wd.inflect({word.tag.case})[0]
                        if i.istitle() is True:
                            wd = wd.capitalize()
                        offers[k] = re.sub(r'\b' + i + r'\b', wd, offers[k], count=1)
                    if word.tag.POS == 'ADJF' and 'Apro' in word.tag:
                        list1 = ["мой", "моего", "моему", "моим", "моём", "моем", "твой", "твоего", "твоему", "твоим",
                                 "твоём", "твоем", "моя", "моей", "мою", "твоя", "твоей", "твою", "моё", "мое",
                                 "мои", "моих", "моим", "моими", "твои", "твоих", "твоим", "твоими"]

                        list2 = ["наш", "нашего", "нашему", "нашим", "нашем", "ваш", "вашего", "вашему", "вашим",
                                 "вашем", "наше", "ваше", "наша", "нашей", "нашу", "ваша", "вашей", "вашу",
                                 "наши", "наших", "нашим", "нашими", "ваши", "ваших", "вашим", "вашими"]
                        for j in list1:
                            if i == j or i == j.capitalize():
                                if gender == 'masc':
                                    wd = 'его'
                                else:
                                    wd = 'её'
                                if i.istitle() is True:
                                    wd = wd.capitalize()
                                offers[k] = re.sub(r'\b' + i + r'\b', wd, offers[k], count=1)
                                break
                        for j in list2:
                            if i == j or i == j.capitalize():
                                wd = 'их'
                                if i.istitle() is True:
                                    wd = wd.capitalize()
                                offers[k] = re.sub(r'\b' + i + r'\b', wd, offers[k], count=1)
                                break
                    if {'VERB', '1per', 'sing', 'pres'} in word.tag:
                        wd = word.inflect({'VERB', 'sing', 'pres', '3per'})[0]
                        if i.istitle() is True:
                            wd = wd.capitalize()
                        offers[k] = re.sub(r'\b' + i + r'\b', wd, offers[k], count=1)
                for i in Dictionary_After1_Error:
                    offers[k] = re.sub(str(i), Dictionary_After1_Error[i], offers[k])
            for i in range(0, len(offers)):
                if offers[i] != '':
                    if offers[i].endswith('!') or offers[i].endswith('?'):
                        textline += offers[i]
                    else:
                        textline += offers[i] + '.'
                else:
                    continue
            my_text2.insert(1.0, textline + '\n')
            Count += 1

        Count = 0
        while Count != 1:
            file = my_text.get(1.0, END).split("\n")
            file2 = my_text2.get(1.0, END).split("\n")
            my_text.delete(1.0, END)
            my_text2.delete(1.0, END)

            for i in range(0, len(file)):
                string1 = file[i].split()
                string2 = file2[i].split()
                offset_1 = ""
                offset_2 = ""
                for j in range(0, len(string1)):
                    if string1[j] == string2[j]:
                        my_text.insert(END, string1[j] + ' ')
                        my_text2.insert(END, string2[j] + ' ')
                        offset_1 += string1[j] + ' '
                        offset_2 += string2[j] + ' '
                    else:
                        if j != len(string1):
                            my_text.insert(END, string1[j] + ' ')
                            my_text2.insert(END, string2[j] + ' ')
                            offset_1 += string1[j] + ' '
                            offset_2 += string2[j] + ' '
                        else:
                            my_text.insert(END, string1[j] + '. ')
                            my_text2.insert(END, string2[j] + '. ')
                            offset_1 += string1[j] + '. '
                            offset_2 += string2[j] + '. '

                        word1 = string1[j]
                        offset1 = '+%dc' % len(word1)
                        word2 = string2[j]
                        offset2 = '+%dc' % len(word2)

                        offset_11 = '+%dc' % (len(offset_1)-1-len(word1))
                        offset_22 = '+%dc' % (len(offset_2)-1-len(word2))
                        pos1 = str(float(i+1))+offset_11
                        pos2 = str(float(i+1))+offset_22

                        pos_start1 = my_text.search(word1, pos1, END)
                        pos_start2 = my_text2.search(word2, pos2, END)

                        pos_end1 = pos_start1 + offset1
                        pos_end2 = pos_start2 + offset2

                        my_text.tag_add("Было", pos_start1, pos_end1)
                        my_text2.tag_add("Стало", pos_start2, pos_end2)
                my_text.insert(END, '\n')
                my_text2.insert(END, '\n')
            Count += 1

def add_file():
    global filepath
    filepath = askopenfilename(initialdir=".", filetypes=(("Docx File", "*.docx"), ("Open Office File", "*.odt"), ("All Files", "*.*")), title="Выберите файл")
    if filepath == "":
        return
    else:
        my_text.delete(1.0, END)
        if filepath.endswith(".docx"):
            file = Document(filepath)
            Str = 52

            Count = 0
            while Count != 2:
                if file.paragraphs[Str].text == '':
                    Count += 1
                    Str += 1
                    continue
                else:
                    my_text.insert(END, file.paragraphs[Str].text + '\n')
                    Str += 1
        else:
            file = load(filepath)
            Str = 55
            allparas = file.getElementsByType(text.P)

            Count = 0
            while Count != 2:
                if teletype.extractText(allparas[Str]) == '':
                    Count += 1
                    Str += 1
                    continue
                else:
                    my_text.insert(END, teletype.extractText(allparas[Str]))
                    Str += 1

def copy_text(event):
    global selected
    if event:
        selected = window.clipboard_get()
    if my_text.selection_get():
        selected = my_text.selection_get()
        window.clipboard_clear()
        window.clipboard_append(selected)

def paste_text(event):
    global selected
    if event:
        selected = window.clipboard_get()
    else:
        if selected:
            position = my_text.index(INSERT)
            my_text.insert(position, selected)

def select_all_text(event):
    my_text.tag_add('sel', '1.0', 'end')

def cut_text(event):
    global selected
    if event:
        selected = window.clipboard_get()
    else:
        if my_text.selection_get():
            selected = my_text.selection_get()
            my_text.delete("sel.first", "sel.last")
            window.clipboard_clear()
            window.clipboard_append(selected)

def clear_text():
    global filepath
    my_text.delete(1.0, END)
    filepath = ""
    sex.set(0)

def clear_text2():
    global filepath
    my_text2.delete(1.0, END)
    filepath = ""
    sex.set(0)

def exit_conv(event):
    #тут надо чекбокс
    answer = mb.askyesno(title="Выход", message="Вы действительно хотите выйти из программы??")
    if answer:
        window.destroy()
    #sys.exit() не работает но пусть пока будет

def save_file():
    name = asksaveasfilename(defaultextension=("Text Files", "*.txt"))
    if name == '':
        return
    else:
        if os.path.split(name)[1].endswith(".docx") or os.path.split(name)[1].endswith(".odt"):
            if os.path.split(name)[1].endswith(".docx"):
                if os.path.isfile(name) is False:
                    document = docx.Document()
                    document.save(os.path.split(name)[1])
                file = Document(name)
                text2save = str(my_text2.get(1.0, END))
                file.add_paragraph(text2save)
            else:
                if os.path.isfile(name) is False:
                    document = OpenDocumentText()
                    document.save(os.path.split(name)[1])
                file = load(name)
                text2save = my_text2.get(1.0, END)
                file.text.addElement(text.P(text=text2save))
            file.save(os.path.split(name)[1])
        else:
            file = open(name, 'w')
            text2save = str(my_text2.get(1.0, END))
            file.write(text2save)
            file.close()


#Полезные константы
filepath = ""

# основные настройки
window = Tk()
window.title("Конвертер")
window.geometry("1200x650")
# радиобокс для выбора гендера
sex = IntVar()  # храним как int, можно как bool, тогда меняем на BooleanVar()
sex.set(0)  # по умолчанию будет мужской ибо value = 0 у мужского рода и в sex.set в скобках 0
male = Radiobutton(text="Мужской род", variable=sex, value=0)

female = Radiobutton(text="Женский род", variable=sex, value=1)

print_file = BooleanVar()
print_file.set(0)
ch_box = Checkbutton(window, text="Выгружать результат в файл", variable=print_file, onvalue=1, offvalue=0)

# меню для файла
menu = Menu(window)
new_item = Menu(menu)
new_item.add_command(label="Загрузить файл", command=add_file)
menu.add_cascade(label='Файл', menu=new_item)
# меню для редактирования строк
new_item = Menu(menu)
new_item.add_command(label="Копировать", command=lambda: copy_text(0), accelerator="(Ctrl+C)")
new_item.add_command(label="Вставить", command=lambda: paste_text(0), accelerator="(Ctrl+V)")
new_item.add_command(label="Вырезать", command=lambda: cut_text(0), accelerator="(Ctrl+X)")
new_item.add_command(label="Выделить всё", command=lambda: select_all_text(0), accelerator="(Ctrl+A)")
menu.add_cascade(label='Редактировать', menu=new_item)
window.config(menu=menu)
# поясняшка
lbl = Label(window, text="Вставьте в поле ввода необходимый фрагмент текста или загрузите файл и нажмите ПРЕОБРАЗОВАТЬ", font="16")
lbl.place(relx=0.12, rely=0.09)

lbl1 = Label(window, text="Поле ввода")
lbl2 = Label(window, text="Поле вывода")
lbl1.place(relx=0.1, rely=0.15)
lbl2.place(relx=0.60, rely=0.15)
# текстовое поле
my_text = Text(window, width=62, height=30)
my_text.pack(pady=5, padx=5)
my_text.place(relx=0.05, rely=0.2)

my_text2 = Text(window, width=62, height=30)
my_text2.pack(pady=5, padx=5)
my_text2.place(relx=0.55, rely=0.2)

my_text.tag_config("Было", background="red", foreground="black")
my_text2.tag_config("Стало", background="yellow", foreground="blue")

# кнопушки
btn_clear = Button(window, text="Очистить поле ввода", command=clear_text)
btn_conv = Button(window, text="ПРЕОБРАЗОВАТЬ", command=conv)
btn_upload = Button(window, text="Загрузить файл", command=add_file)
btn_clear2 = Button(window, text="Очистить поле вывода", command=clear_text2)
btn_save = Button(window, text="Сохранить файл", command=save_file)

#выравнивание кнопушек
btn_clear.grid(row=0, column=0)
btn_clear2.grid(row=0, column=1)
btn_upload.grid(row=0, column=2)
btn_conv.grid(row=0, column=3)
btn_save.grid(row=0, column=4)
male.grid(row=0, column=5)
female.grid(row=0, column=6)
ch_box.grid(row=0, column=7)


#горячие клавиши
window.bind("<Control-Key-c>", copy_text)
window.bind("<Control-Key-V>", paste_text)
window.bind("<Control-Key-X>", cut_text)
window.bind("<Control-Key-A>", select_all_text)
window.bind("<Escape>", exit_conv)

window.mainloop()
